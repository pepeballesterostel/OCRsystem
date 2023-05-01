import os
import pytesseract
import docx
import argparse
import requests
from googletrans import Translator
from tqdm import tqdm
import utils


def main(args):
    # Check if the image database path exists
    assert os.path.exists(args.database), "Image database path does not exist"

    # Check if the Tesseract OCR engine is installed
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    try:
        pytesseract.get_tesseract_version()
    except pytesseract.TesseractNotFoundError:
        assert False, "Tesseract OCR engine is not installed"

    # Check if the Tesseract OCR engine path is correct
    assert os.path.exists(pytesseract.pytesseract.tesseract_cmd), "Invalid Tesseract OCR engine path"

    # Set the TESSDATA_PREFIX environment variable
    os.environ['TESSDATA_PREFIX'] = './tessdata/'
    # Create the tessdata directory if it doesn't exist
    if not os.path.exists('tessdata'):
        os.makedirs('tessdata')

    # Download the language data file if it doesn't exist
    data_file_path = f"tessdata/{args.language}.traineddata"
    if not os.path.exists(data_file_path):
        print(f"Downloading {args.language}.traineddata file...")
        url = f"https://github.com/tesseract-ocr/tessdata_best/raw/main/{args.language}.traineddata"
        response = requests.get(url)
        with open(data_file_path, "wb") as f:
            f.write(response.content)
        print(f"{args.language}.traineddata file downloaded successfully.")
    else:
        print(f"{args.language}.traineddata file already exists.")

    # Create a new Word document
    doc = docx.Document()
    doc.add_heading('Transtaled text from: {}'.format(str(args.database)), 0)

    translator = Translator()  

    # Iterate through each image in the directory and convert it to text
    for filename in tqdm(os.listdir(args.database)):
        if filename.endswith(".jpg") or filename.endswith(".png") or filename.endswith(".JPG") or filename.endswith(".PNG") or filename.endswith(".jpeg") or filename.endswith(".JPEG"):
            print('\n')
            print(f"Processing file: {filename}")
            try:
                filename = os.path.join(args.database, filename)
                img = utils.process_image(filename)
                text = pytesseract.image_to_string(img, lang = args.language, config='--psm 6')
                if args.target_lang != '':
                    try:
                        text = utils.process_text(text)
                        text = translator.translate(text, dest=args.target_lang).text
                    except ValueError:
                        print(f"Could not translate text from file: {filename}")
                        pass
                doc.add_paragraph(text)
                utils.add_delimiter(doc)
            except FileNotFoundError:
                print(f"Could not find file: {filename}")
            except TypeError as e:
                print(f"Error processing file: {filename}. Reason: {e}")

    # check if the output path exists
    if os.path.exists(args.output):
        print(f"File {args.output} already exists. Do you want to overwrite it? [y/n]")
        answer = input()
        if answer == 'y':
            doc.save(args.output)
        else:
            print("File not saved.")
    else:
        doc.save(args.output)

if __name__ == '__main__':
    # Create an argument parser
    parser = argparse.ArgumentParser(description='Extract text from book images')
    parser.add_argument('-d', '--database', default='Baxandall', type=str,
                        help='Path to the book image database')
    parser.add_argument('-l', '--language', default='eng', type=str,
                        help='Language of the book text (e.g. "deu" for German, "ita" for Italian). Take into account that the data for each language must be downloaded from https://github.com/tesseract-ocr/tessdata')
    parser.add_argument('-o', '--output', default='extracted_text.docx',
                        help='Path to the output Word document')
    parser.add_argument('-tl', '--target_lang', default='es', type=str, 
                        help='Language to translate text to (e.g. "en" for English, "es" for Spanish). If not specified, the text will not be translated.')


    # Parse the command-line arguments
    args = parser.parse_args()
    args = parser.parse_args()
    main(args)